# encoding: utf-8

"""
Test suite for the docx.text module
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.text import Paragraph, Run

import pytest

from .unitutil.cxml import element, xml
from .unitutil.mock import call, class_mock, instance_mock


class DescribeParagraph(object):

    def it_provides_access_to_the_runs_it_contains(self, runs_fixture):
        paragraph, Run_, r_, r_2_, run_, run_2_ = runs_fixture
        runs = paragraph.runs
        assert Run_.mock_calls == [
            call(r_, paragraph), call(r_2_, paragraph)
        ]
        assert runs == [run_, run_2_]

    def it_can_add_a_run_to_itself(self, add_run_fixture):
        paragraph, text, style, expected_xml = add_run_fixture
        run = paragraph.add_run(text, style)
        assert paragraph._p.xml == expected_xml
        assert isinstance(run, Run)
        assert run._r is paragraph._p.r_lst[0]

    def it_knows_its_alignment_value(self, alignment_get_fixture):
        paragraph, expected_value = alignment_get_fixture
        assert paragraph.alignment == expected_value

    def it_can_change_its_alignment_value(self, alignment_set_fixture):
        paragraph, value, expected_xml = alignment_set_fixture
        paragraph.alignment = value
        assert paragraph._p.xml == expected_xml

    def it_knows_its_paragraph_style(self, style_get_fixture):
        paragraph, expected_style = style_get_fixture
        assert paragraph.style == expected_style

    def it_can_change_its_paragraph_style(self, style_set_fixture):
        paragraph, value, expected_xml = style_set_fixture
        paragraph.style = value
        assert paragraph._p.xml == expected_xml

    def it_knows_the_text_it_contains(self, text_get_fixture):
        paragraph, expected_text = text_get_fixture
        assert paragraph.text == expected_text

    def it_can_replace_the_text_it_contains(self, text_set_fixture):
        paragraph, text, expected_text = text_set_fixture
        paragraph.text = text
        assert paragraph.text == expected_text

    def it_can_insert_a_paragraph_before_itself(self, insert_before_fixture):
        paragraph, text, style, body, expected_xml = insert_before_fixture
        new_paragraph = paragraph.insert_paragraph_before(text, style)
        assert isinstance(new_paragraph, Paragraph)
        assert new_paragraph.text == text
        assert new_paragraph.style == style
        assert body.xml == expected_xml

    def it_can_remove_its_content_while_preserving_formatting(
            self, clear_fixture):
        paragraph, expected_xml = clear_fixture
        _paragraph = paragraph.clear()
        assert paragraph._p.xml == expected_xml
        assert _paragraph is paragraph

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        ('w:p', None, None,
         'w:p/w:r'),
        ('w:p', 'foobar', None,
         'w:p/w:r/w:t"foobar"'),
        ('w:p', None, 'Strong',
         'w:p/w:r/w:rPr/w:rStyle{w:val=Strong}'),
        ('w:p', 'foobar', 'Strong',
         'w:p/w:r/(w:rPr/w:rStyle{w:val=Strong}, w:t"foobar")'),
    ])
    def add_run_fixture(self, request):
        before_cxml, text, style, after_cxml = request.param
        paragraph = Paragraph(element(before_cxml), None)
        expected_xml = xml(after_cxml)
        return paragraph, text, style, expected_xml

    @pytest.fixture(params=[
        ('w:p/w:pPr/w:jc{w:val=center}', WD_ALIGN_PARAGRAPH.CENTER),
        ('w:p', None),
    ])
    def alignment_get_fixture(self, request):
        cxml, expected_alignment_value = request.param
        paragraph = Paragraph(element(cxml), None)
        return paragraph, expected_alignment_value

    @pytest.fixture(params=[
        ('w:p', WD_ALIGN_PARAGRAPH.LEFT,
         'w:p/w:pPr/w:jc{w:val=left}'),
        ('w:p/w:pPr/w:jc{w:val=left}', WD_ALIGN_PARAGRAPH.CENTER,
         'w:p/w:pPr/w:jc{w:val=center}'),
        ('w:p/w:pPr/w:jc{w:val=left}', None,
         'w:p/w:pPr'),
        ('w:p', None, 'w:p/w:pPr'),
    ])
    def alignment_set_fixture(self, request):
        initial_cxml, new_alignment_value, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_alignment_value, expected_xml

    @pytest.fixture(params=[
        ('w:p', 'w:p'),
        ('w:p/w:pPr', 'w:p/w:pPr'),
        ('w:p/w:r/w:t"foobar"', 'w:p'),
        ('w:p/(w:pPr, w:r/w:t"foobar")', 'w:p/w:pPr'),
    ])
    def clear_fixture(self, request):
        initial_cxml, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, expected_xml

    @pytest.fixture(params=[
        ('w:body/w:p', 'foobar', 'Heading1',
         'w:body/(w:p/(w:pPr/w:pStyle{w:val=Heading1},w:r/w:t"foobar"),w:p)')
    ])
    def insert_before_fixture(self, request):
        body_cxml, text, style, expected_cxml = request.param
        body = element(body_cxml)
        paragraph = Paragraph(body.find(qn('w:p')), None)
        expected_xml = xml(expected_cxml)
        return paragraph, text, style, body, expected_xml

    @pytest.fixture
    def runs_fixture(self, p_, Run_, r_, r_2_, runs_):
        paragraph = Paragraph(p_, None)
        run_, run_2_ = runs_
        return paragraph, Run_, r_, r_2_, run_, run_2_

    @pytest.fixture(params=[
        ('w:p', 'Normal'),
        ('w:p/w:pPr', 'Normal'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', 'Heading1'),
    ])
    def style_get_fixture(self, request):
        p_cxml, expected_style = request.param
        paragraph = Paragraph(element(p_cxml), None)
        return paragraph, expected_style

    @pytest.fixture(params=[
        ('w:p',                                'Heading1',
         'w:p/w:pPr/w:pStyle{w:val=Heading1}'),
        ('w:p/w:pPr',                          'Heading1',
         'w:p/w:pPr/w:pStyle{w:val=Heading1}'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', 'Heading2',
         'w:p/w:pPr/w:pStyle{w:val=Heading2}'),
        ('w:p/w:pPr/w:pStyle{w:val=Heading1}', None,
         'w:p/w:pPr'),
        ('w:p',                                None,
         'w:p/w:pPr'),
    ])
    def style_set_fixture(self, request):
        p_cxml, new_style_value, expected_cxml = request.param
        paragraph = Paragraph(element(p_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_style_value, expected_xml

    @pytest.fixture(params=[
        ('w:p', ''),
        ('w:p/w:r', ''),
        ('w:p/w:r/w:t', ''),
        ('w:p/w:r/w:t"foo"', 'foo'),
        ('w:p/w:r/(w:t"foo", w:t"bar")', 'foobar'),
        ('w:p/w:r/(w:t"fo ", w:t"bar")', 'fo bar'),
        ('w:p/w:r/(w:t"foo", w:tab, w:t"bar")', 'foo\tbar'),
        ('w:p/w:r/(w:t"foo", w:br,  w:t"bar")', 'foo\nbar'),
        ('w:p/w:r/(w:t"foo", w:cr,  w:t"bar")', 'foo\nbar'),
    ])
    def text_get_fixture(self, request):
        p_cxml, expected_text_value = request.param
        paragraph = Paragraph(element(p_cxml), None)
        return paragraph, expected_text_value

    @pytest.fixture
    def text_set_fixture(self):
        paragraph = Paragraph(element('w:p'), None)
        paragraph.add_run('must not appear in result')
        new_text_value = 'foo\tbar\rbaz\n'
        expected_text_value = 'foo\tbar\nbaz\n'
        return paragraph, new_text_value, expected_text_value

    # fixture components ---------------------------------------------

    @pytest.fixture
    def p_(self, request, r_, r_2_):
        return instance_mock(request, CT_P, r_lst=(r_, r_2_))

    @pytest.fixture
    def Run_(self, request, runs_):
        run_, run_2_ = runs_
        return class_mock(
            request, 'docx.text.Run', side_effect=[run_, run_2_]
        )

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def r_2_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def runs_(self, request):
        run_ = instance_mock(request, Run, name='run_')
        run_2_ = instance_mock(request, Run, name='run_2_')
        return run_, run_2_
